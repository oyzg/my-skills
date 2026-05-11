#!/usr/bin/env python3
"""
文章导出工具（支持图片插入与自动裁剪四周）
支持导出为 Word (.docx)、PDF、Markdown (.md) 格式
"""

import argparse
import os
import re
from typing import Optional, List, Dict, Tuple


def parse_markdown_table(lines: List[str], start_idx: int) -> Tuple[List[List[str]], int]:
    """
    解析Markdown表格

    Args:
        lines: 所有行
        start_idx: 表格起始行索引

    Returns:
        (表格数据矩阵, 表格结束后的行索引)
        表格数据：每行是一个列表，包含各列内容
    """
    table_data = []
    i = start_idx

    # 收集表格行
    while i < len(lines):
        line = lines[i].strip()
        # 检查是否是表格行（包含 |）
        if '|' not in line:
            break

        # 分割列
        cells = [cell.strip() for cell in line.split('|')]

        # 移除首尾空列（由于 | 开头和结尾导致）
        if cells and cells[0] == '':
            cells = cells[1:]
        if cells and cells[-1] == '':
            cells = cells[:-1]

        # 检查是否是分隔行（只包含 - 和 |）
        if all(cell.replace('-', '').replace(':', '').strip() == '' for cell in cells):
            i += 1
            continue

        table_data.append(cells)
        i += 1

    return table_data, i


def set_table_style(table):
    """
    设置表格样式（学术论文风格）

    Args:
        table: Word表格对象
    """
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # 设置表格宽度（自动）
    table.autofit = True

    # 遍历所有行和单元格
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            # 设置单元格内边距（适中的留白）
            cell.margin_top = Pt(3)
            cell.margin_bottom = Pt(3)
            cell.margin_left = Pt(6)
            cell.margin_right = Pt(6)

            # 设置单元格边框（细黑色实线）
            set_cell_border(cell,
                          top={"sz": "4", "val": "single", "color": "000000"},  # 0.5pt
                          bottom={"sz": "4", "val": "single", "color": "000000"},
                          left={"sz": "4", "val": "single", "color": "000000"},
                          right={"sz": "4", "val": "single", "color": "000000"})

                # 设置单元格字体
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '微软雅黑')
                    run.font.size = Pt(10.5)  # 约 11pt
                    run.font.color.rgb = RGBColor(0, 0, 0)

                # 表头行样式
                if i == 0:
                    # 设置背景色（浅灰色）
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), 'E0E0E0')  # 稍深一点的灰色，更清晰
                    cell._element.get_or_add_tcPr().append(shading_elm)

                    # 加粗
                    for run in paragraph.runs:
                        run.font.bold = True

                    # 左对齐（不是居中）
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    # 数据行：左对齐
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def set_cell_border(cell, **kwargs):
    """
    设置单元格边框

    Args:
        cell: 单元格对象
        **kwargs: 边框参数
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # 定义边框参数
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # 设置各个边框
    for edge in ('left', 'top', 'right', 'bottom', 'start', 'end'):
        if edge in kwargs:
            edge_data = kwargs[edge]
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def validate_and_fix_markdown(content: str) -> Tuple[str, List[str]]:
    """
    检查并修复Markdown格式问题

    Args:
        content: Markdown内容

    Returns:
        (修复后的内容, 问题列表)
    """
    issues = []
    fixed_content = content
    lines = content.split('\n')

    # 逐行检查和修复
    for i, line in enumerate(lines):
        original_line = line

        # 1. 检查加粗标记配对
        bold_markers = line.count('**')
        if bold_markers % 2 != 0:
            issues.append(f"第{i+1}行：加粗标记不匹配（奇数个**），已自动修复")
            # 找到最后一个**，删除它
            last_idx = line.rfind('**')
            if last_idx != -1:
                line = line[:last_idx] + line[last_idx+2:]

        # 2. 检测并修复连续加粗 **text1****text2**
        if '****' in line and '**' in line:
            # 确保不是空加粗
            if re.search(r'\*\*[^*]+\*\*\*\*[^*]+\*\*', line):
                issues.append(f"第{i+1}行：检测到连续加粗，已添加空格分隔")
                line = re.sub(r'\*\*([^*]+)\*\*\*\*([^*]+)\*\*', r'**\1** **\2**', line)

        # 3. 检测并修复空加粗 ****** 或 ****
        if re.search(r'\*\*\*\*\*\*', line):
            issues.append(f"第{i+1}行：检测到空加粗标记，已删除")
            line = re.sub(r'\*\*\*\*\*\*', '', line)

        # 4. 检测未闭合的加粗（段落末尾有未闭合的**）
        if '**' in line:
            # 计算当前行的加粗状态
            bold_count = line.count('**')
            if bold_count % 2 != 0:
                # 这行有未闭合的加粗
                # 检查是否真的未闭合（不是被空格分隔的）
                parts = line.split('**')
                if len(parts) % 2 != 0:
                    issues.append(f"第{i+1}行：检测到未闭合的加粗，已在段落末尾添加闭合标记")
                    line += '**'

        # 5. 检测嵌套加粗 **text **inner** text**
        if re.search(r'\*\*[^*]*\*\*[^*]*\*\*', line):
            # 检查是否有真正的嵌套
            # 找到第一个**
            first_bold = line.find('**')
            if first_bold != -1:
                # 找到第一个闭合**
                first_close = line.find('**', first_bold + 2)
                if first_close != -1:
                    # 检查第二个加粗是否在第一个闭合之前
                    second_bold = line.find('**', first_close + 2)
                    if second_bold != -1:
                        # 检查第二个闭合
                        second_close = line.find('**', second_bold + 2)
                        if second_close != -1:
                            # 这是嵌套加粗
                            issues.append(f"第{i+1}行：警告：检测到嵌套加粗，可能导致格式错误")
                            # 不自动修复，只是警告

        if line != original_line:
            fixed_content = fixed_content.replace(original_line, line)

    # 6. 检查斜体标记配对
    italic_lines = []
    for i, line in enumerate(fixed_content.split('\n')):
        # 排除加粗标记
        clean_line = line.replace('**', '')
        italic_count = clean_line.count('*')
        if italic_count % 2 != 0:
            issues.append(f"第{i+1}行：警告：斜体标记可能不匹配（排除加粗后仍有奇数个*）")

    # 7. 检查代码标记配对
    code_lines = []
    for i, line in enumerate(fixed_content.split('\n')):
        code_count = line.count('`')
        if code_count % 2 != 0:
            issues.append(f"第{i+1}行：警告：代码标记可能不匹配（奇数个`）")

    return fixed_content, issues


def optimize_word_document_styles(doc):
    """
    优化Word文档的排版样式

    Args:
        doc: Word文档对象
    """
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_LINE_SPACING

    # 1. 优化普通段落样式
    normal_style = doc.styles['Normal']
    normal_style.font.name = '宋体'
    normal_style._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '宋体')
    normal_style.font.size = Pt(12)
    # 设置段落间距
    normal_style.paragraph_format.space_before = Pt(6)
    normal_style.paragraph_format.space_after = Pt(6)
    normal_style.paragraph_format.line_spacing = 1.5  # 1.5倍行距

    # 2. 优化标题样式
    for level in [1, 2, 3]:
        heading_style = doc.styles[f'Heading {level}']
        heading_style.paragraph_format.space_before = Pt(12)
        heading_style.paragraph_format.space_after = Pt(6)
        heading_style.paragraph_format.line_spacing = 1.5

    # 3. 优化列表样式
    if 'List Bullet' in doc.styles:
        list_bullet = doc.styles['List Bullet']
        list_bullet.font.name = '宋体'
        list_bullet.font.size = Pt(12)
        list_bullet.paragraph_format.space_before = Pt(3)
        list_bullet.paragraph_format.space_after = Pt(3)
        list_bullet.paragraph_format.left_indent = Inches(0.25)

    if 'List Number' in doc.styles:
        list_number = doc.styles['List Number']
        list_number.font.name = '宋体'
        list_number.font.size = Pt(12)
        list_number.paragraph_format.space_before = Pt(3)
        list_number.paragraph_format.space_after = Pt(3)
        list_number.paragraph_format.left_indent = Inches(0.25)



def crop_borders(image_path: str, output_path: Optional[str] = None, crop_ratio: float = 0.06) -> bool:
    """
    裁剪图片四周（去除水印）

    Args:
        image_path: 输入图片路径
        output_path: 输出图片路径，如果为 None 则覆盖原文件
        crop_ratio: 裁剪比例（每侧裁剪的比例，默认 0.06 即 6%）

    Returns:
        是否成功裁剪
    """
    try:
        from PIL import Image
    except ImportError:
        print("警告：缺少 Pillow 库，无法裁剪图片。请安装：pip install Pillow")
        return False

    if output_path is None:
        output_path = image_path

    try:
        # 打开图片
        img = Image.open(image_path)
        width, height = img.size

        # 计算裁剪区域（每侧裁剪 crop_ratio）
        crop_left = int(width * crop_ratio)
        crop_top = int(height * crop_ratio)
        crop_right = int(width * (1 - crop_ratio))
        crop_bottom = int(height * (1 - crop_ratio))

        # 检查裁剪后尺寸是否合理
        if crop_left >= crop_right or crop_top >= crop_bottom:
            print(f"裁剪比例过大，跳过裁剪：{image_path}")
            return False

        # 裁剪图片（左上右下）
        cropped = img.crop((crop_left, crop_top, crop_right, crop_bottom))
        cropped.save(output_path)
        print(f"已裁剪四周：{image_path} -> {output_path} (裁剪比例: {crop_ratio*100:.0f}%)")
        return True

    except Exception as e:
        print(f"裁剪失败：{image_path}，错误：{e}")
        return False


def process_images_with_crop(images: List[Dict], crop_borders: bool = True) -> List[Dict]:
    """
    处理图片列表，自动裁剪四周

    Args:
        images: 图片信息列表
        crop_borders: 是否裁剪四周

    Returns:
        处理后的图片信息列表
    """
    if not crop_borders:
        return images

    processed_images = []
    for image in images:
        image_path = image['image_path']

        # 检查图片是否存在
        if os.path.exists(image_path):
            # 裁剪四周
            success = crop_borders(image_path, crop_ratio=0.06)

            # 记录处理结果
            if success:
                image['borders_cropped'] = True
            else:
                image['borders_cropped'] = False

            processed_images.append(image)
        else:
            print(f"警告：图片文件不存在 {image_path}")
            image['borders_cropped'] = False
            processed_images.append(image)

    return processed_images


def parse_images_from_content(article_content: str, images_dir: str = "./images/") -> List[Dict]:
    """
    从文章内容中解析图片标记

    Args:
        article_content: 文章内容
        images_dir: 图片目录

    Returns:
        图片信息列表，每个元素包含：position, description, image_path
    """
    images = []
    lines = article_content.split('\n')

    # 匹配 [IMAGE: 描述] 格式
    pattern = re.compile(r'\[IMAGE:\s*(.*?)\]')

    for i, line in enumerate(lines):
        match = pattern.search(line)
        if match:
            description = match.group(1)
            # 尝试提取图片编号（image_01, image_02 等）
            image_num_match = re.search(r'image_(\d+)', description, re.IGNORECASE)
            if image_num_match:
                image_num = image_num_match.group(1)
                image_filename = f"image_{image_num.zfill(2)}.jpg"
                image_path = os.path.join(images_dir, image_filename)
            else:
                # 如果没有编号，使用默认编号
                image_path = os.path.join(images_dir, f"image_{len(images)+1:02d}.jpg")

            images.append({
                'line_number': i,
                'description': description,
                'image_path': image_path,
                'line_content': line
            })

    return images


def export_to_markdown(article_content: str, output_filename: str, images_dir: str = "./images/",
                       crop_borders: bool = True) -> str:
    """
    导出文章为 Markdown 格式（支持图片插入）

    Args:
        article_content: 文章内容
        output_filename: 输出文件名（不含扩展名）
        images_dir: 图片目录
        crop_borders: 是否裁剪图片四周

    Returns:
        生成的文件路径
    """
    filename = f"{output_filename}.md"
    filepath = os.path.join(".", filename)

    # 解析图片
    images = parse_images_from_content(article_content, images_dir)

    # 处理图片
    if crop_borders:
        images = process_images_with_crop(images, crop_borders)

    # 替换图片标记
    lines = article_content.split('\n')
    result_lines = []

    for i, line in enumerate(lines):
        # 检查是否是图片标记行
        is_image_line = False
        for image in images:
            if image['line_number'] == i:
                is_image_line = True
                # 插入图片 Markdown 语法
                image_relative_path = os.path.relpath(image['image_path'], ".")
                img_markdown = "\n\n![{}]({})\n\n".format(image['description'], image_relative_path)
                result_lines.append(img_markdown)
                break

        if not is_image_line:
            result_lines.append(line)

    # 写入文件
    with open(filepath, "w", encoding="utf-8") as f:
        f.write('\n'.join(result_lines))

    return filepath


def parse_markdown_to_runs(text: str, paragraph):
    """
    解析 Markdown 行内格式并添加到段落

    Args:
        text: 包含 Markdown 格式的文本
        paragraph: Word 段落对象

    Returns:
        None（直接修改段落）
    """
    from docx.shared import RGBColor

    # 定义格式模式及其处理函数
    patterns = [
        # 行内代码 `code`
        (r'`([^`]+)`', lambda m: (m.group(1), {'code': True})),
        # 链接 [text](url)
        (r'\[([^\]]+)\]\(([^)]+)\)', lambda m: (m.group(1), {'link': m.group(2)})),
        # 加粗 **text**
        (r'\*\*([^*]+)\*\*', lambda m: (m.group(1), {'bold': True})),
        # 斜体 *text*（注意要排除 ** 已匹配的情况）
        (r'(?<!\*)\*([^*]+)\*(?!\*)', lambda m: (m.group(1), {'italic': True})),
        # 删除线 ~~text~~
        (r'~~([^~]+)~~', lambda m: (m.group(1), {'strike': True})),
    ]

    # 分割文本为片段
    segments = [(text, {})]

    for pattern, handler in patterns:
        new_segments = []
        for seg_text, seg_style in segments:
            # 跳过已经是代码格式的段落（代码不进行嵌套解析）
            if seg_style.get('code'):
                new_segments.append((seg_text, seg_style))
                continue

            # 查找所有匹配
            import re
            matches = list(re.finditer(pattern, seg_text))
            if not matches:
                new_segments.append((seg_text, seg_style))
                continue

            # 分割文本
            last_end = 0
            for match in matches:
                # 添加匹配前的普通文本
                if match.start() > last_end:
                    normal_text = seg_text[last_end:match.start()]
                    if normal_text:
                        new_segments.append((normal_text, seg_style.copy()))

                # 添加格式化文本
                formatted_text, style_props = handler(match)
                merged_style = seg_style.copy()
                merged_style.update(style_props)
                new_segments.append((formatted_text, merged_style))

                last_end = match.end()

            # 添加剩余文本
            if last_end < len(seg_text):
                remaining_text = seg_text[last_end:]
                if remaining_text:
                    new_segments.append((remaining_text, seg_style.copy()))

        segments = new_segments

    # 将所有片段添加到段落
    for seg_text, seg_style in segments:
        run = paragraph.add_run(seg_text)

        # 应用样式
        if seg_style.get('bold'):
            run.font.bold = True
        if seg_style.get('italic'):
            run.font.italic = True
        if seg_style.get('strike'):
            run.font.strike = True
        if seg_style.get('code'):
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(200, 50, 50)
        if seg_style.get('link'):
            # 添加超链接
            hyperlink_url = seg_style['link']
            # 注意：docx 的超链接需要特殊处理，这里简化处理
            # 实际应用中可能需要使用 docx.oxml 来添加真正的超链接
            pass


def export_to_word(article_content: str, output_filename: str, images_dir: str = "./images/",
                   crop_borders: bool = True) -> Optional[str]:
    """
    导出文章为 Word 格式（支持图片插入）

    Args:
        article_content: 文章内容（Markdown 格式）
        output_filename: 输出文件名（不含扩展名）
        images_dir: 图片目录
        crop_borders: 是否裁剪图片四周

    Returns:
        生成的文件路径，如果失败则返回 None
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("错误：缺少 python-docx 库，请安装：pip install python-docx")
        return None

    filename = f"{output_filename}.docx"
    filepath = os.path.join(".", filename)

    # 创建 Word 文档
    doc = Document()

    # 优化文档排版样式
    optimize_word_document_styles(doc)

    # 格式检查和修复
    print("正在检查和修复Markdown格式...")
    fixed_content, format_issues = validate_and_fix_markdown(article_content)
    if format_issues:
        print(f"发现并修复 {len(format_issues)} 个格式问题：")
        for issue in format_issues[:5]:  # 只显示前5个
            print(f"  - {issue}")
        if len(format_issues) > 5:
            print(f"  ... 还有 {len(format_issues) - 5} 个问题")
    else:
        print("✓ 格式检查通过，未发现问题")

    # 解析图片
    images = parse_images_from_content(fixed_content, images_dir)

    # 处理图片
    if crop_borders:
        images = process_images_with_crop(images, crop_borders)

    image_map = {img['line_number']: img for img in images}

    # 解析 Markdown 内容并写入 Word
    lines = fixed_content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
            continue

        # 检查是否是图片标记行
        if i in image_map:
            image = image_map[i]
            # 插入图片
            if os.path.exists(image['image_path']):
                try:
                    # 调整图片大小（宽度 6 英寸）
                    doc.add_picture(image['image_path'], width=Inches(6))
                    # 居中对齐
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # 添加图片说明
                    p = doc.add_paragraph(image['description'])
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.runs[0].font.size = Pt(10)
                    p.runs[0].font.color.rgb = RGBColor(128, 128, 128)
                except Exception as e:
                    print(f"警告：插入图片失败 {image['image_path']}：{e}")
            else:
                print(f"警告：图片文件不存在 {image['image_path']}")
            continue

        # 处理标题
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.runs[0].font.size = Pt(18)
            p.runs[0].font.bold = True

        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.runs[0].font.size = Pt(16)
            p.runs[0].font.bold = True

        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.runs[0].font.size = Pt(14)
            p.runs[0].font.bold = True

        # 处理无序列表
        elif line.startswith('- ') or line.startswith('* '):
            # 提取列表项文本（去掉标记）
            item_text = line[2:] if line.startswith('- ') else line[2:]
            p = doc.add_paragraph(item_text, style='List Bullet')
            # 解析列表项中的行内格式
            parse_markdown_to_runs(item_text, p)

        # 处理有序列表（数字. 格式）
        elif re.match(r'^\d+\.\s', line):
            # 提取列表项文本（去掉数字标记）
            match = re.match(r'^\d+\.\s(.*)$', line)
            item_text = match.group(1) if match else line
            p = doc.add_paragraph(item_text, style='List Number')
            # 解析列表项中的行内格式
            parse_markdown_to_runs(item_text, p)

        # 处理引用
        elif line.startswith('> '):
            # 提取引用文本（去掉 > ）
            quote_text = line[2:]
            p = doc.add_paragraph(quote_text)
            # 设置引用样式
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent = Inches(0.5)
            p.runs[0].font.color.rgb = RGBColor(100, 100, 100)
            p.runs[0].font.italic = True
            # 解析引用中的行内格式
            parse_markdown_to_runs(quote_text, p)

        # 处理分割线
        elif line == '---' or line == '***':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run('—' * 30)
            run.font.color.rgb = RGBColor(200, 200, 200)

        # 处理Markdown表格
        elif '|' in line:
            # 尝试解析表格
            table_data, next_i = parse_markdown_table(lines, i)
            if len(table_data) >= 2:  # 至少有表头和数据
                # 创建表格
                rows = len(table_data)
                cols = len(table_data[0]) if table_data else 0
                table = doc.add_table(rows=rows, cols=cols)

                # 填充表格数据
                for row_idx, row_data in enumerate(table_data):
                    for col_idx, cell_text in enumerate(row_data):
                        if col_idx < len(table.rows[row_idx].cells):
                            cell = table.rows[row_idx].cells[col_idx]
                            cell.text = cell_text

                # 设置表格样式
                set_table_style(table)

                # 跳过已处理的行
                i = next_i - 1
            else:
                # 不是有效的表格，按普通段落处理
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                parse_markdown_to_runs(line, p)

        # 处理代码块（``` 开始）
        elif line.startswith('```'):
            # 跳过代码块标记本身
            pass

        # 处理普通段落
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            # 解析 Markdown 行内格式
            parse_markdown_to_runs(line, p)

        i += 1

    # 保存文档
    doc.save(filepath)

    return filepath


def export_to_pdf(article_content: str, output_filename: str, images_dir: str = "./images/",
                  crop_borders: bool = True) -> Optional[str]:
    """
    导出文章为 PDF 格式（支持图片插入）

    Args:
        article_content: 文章内容（Markdown 格式）
        output_filename: 输出文件名（不含扩展名）
        images_dir: 图片目录
        crop_borders: 是否裁剪图片四周

    Returns:
        生成的文件路径，如果失败则返回 None
    """
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
        from reportlab.lib.enums import TA_LEFT, TA_CENTER
        from reportlab.lib import colors
    except ImportError:
        print("错误：缺少 reportlab 库，请安装：pip install reportlab")
        return None

    filename = f"{output_filename}.pdf"
    filepath = os.path.join(".", filename)

    # 创建 PDF 文档
    doc = SimpleDocTemplate(
        filepath,
        pagesize=A4,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=18
    )

    # 设置样式
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name='CustomHeading',
        parent=styles['Heading1'],
        fontSize=18,
        textColor=colors.HexColor('#000000'),
        spaceAfter=12,
    ))
    styles.add(ParagraphStyle(
        name='CustomBody',
        parent=styles['BodyText'],
        fontSize=12,
        textColor=colors.HexColor('#333333'),
        leading=16,
        spaceAfter=8,
        alignment=TA_LEFT,
    ))
    styles.add(ParagraphStyle(
        name='ImageCaption',
        parent=styles['BodyText'],
        fontSize=10,
        textColor=colors.HexColor('#666666'),
        alignment=TA_CENTER,
        spaceAfter=12,
    ))

    # 解析图片
    images = parse_images_from_content(article_content, images_dir)

    # 处理图片
    if crop_borders:
        images = process_images_with_crop(images, crop_borders)

    image_map = {img['line_number']: img for img in images}

    # 构建 PDF 内容
    story = []
    lines = article_content.split('\n')

    for i, line in enumerate(lines):
        line = line.strip()

        if not line:
            story.append(Spacer(1, 0.2 * inch))
            continue

        # 检查是否是图片标记行
        if i in image_map:
            image = image_map[i]
            # 插入图片
            if os.path.exists(image['image_path']):
                try:
                    img = Image(image['image_path'], width=6 * inch)
                    img.hAlign = 'CENTER'
                    story.append(img)
                    # 添加图片说明
                    caption = Paragraph(image['description'], styles['ImageCaption'])
                    story.append(caption)
                except Exception as e:
                    print(f"警告：插入图片失败 {image['image_path']}：{e}")
            else:
                print(f"警告：图片文件不存在 {image['image_path']}")
            continue

        # 处理标题
        if line.startswith('# '):
            heading = Paragraph(line[2:], styles['CustomHeading'])
            story.append(heading)
        elif line.startswith('## '):
            heading = Paragraph(line[3:], styles['Heading2'])
            story.append(heading)
        elif line.startswith('### '):
            heading = Paragraph(line[4:], styles['Heading3'])
            story.append(heading)
        else:
            # 处理普通段落
            para = Paragraph(line, styles['CustomBody'])
            story.append(para)

    # 生成 PDF
    doc.build(story)

    return filepath


def export_document(article_content: str, format: str, filename: str,
                   images_dir: str = "./images/", crop_borders: bool = True) -> List[str]:
    """
    导出文章

    Args:
        article_content: 文章内容
        format: 导出格式（md/docx/pdf/all）
        filename: 输出文件名（不含扩展名）
        images_dir: 图片目录
        crop_borders: 是否裁剪图片四周

    Returns:
        生成的文件路径列表
    """
    generated_files = []

    if format.lower() in ['md', 'all']:
        filepath = export_to_markdown(article_content, filename, images_dir, crop_borders)
        if filepath:
            generated_files.append(filepath)
            print(f"✓ 已生成 Markdown: {filepath}")

    if format.lower() in ['docx', 'all']:
        filepath = export_to_word(article_content, filename, images_dir, crop_borders)
        if filepath:
            generated_files.append(filepath)
            print(f"✓ 已生成 Word: {filepath}")

    if format.lower() in ['pdf', 'all']:
        filepath = export_to_pdf(article_content, filename, images_dir, crop_borders)
        if filepath:
            generated_files.append(filepath)
            print(f"✓ 已生成 PDF: {filepath}")

    return generated_files


def main():
    parser = argparse.ArgumentParser(description='文章导出工具（支持图片插入与自动裁剪四周）')
    parser.add_argument('--content', type=str, required=True, help='文章内容（支持 Markdown 格式）')
    parser.add_argument('--format', type=str, default='md', choices=['md', 'docx', 'pdf', 'all'],
                        help='导出格式')
    parser.add_argument('--filename', type=str, required=True, help='输出文件名（不含扩展名）')
    parser.add_argument('--images-dir', type=str, default='./images/', help='图片目录')
    parser.add_argument('--crop-borders', action='store_true', default=True,
                        help='是否自动裁剪图片四周（默认启用）')

    args = parser.parse_args()

    # 导出文档
    generated_files = export_document(
        args.content,
        args.format,
        args.filename,
        args.images_dir,
        args.crop_borders
    )

    if generated_files:
        print(f"\n成功生成 {len(generated_files)} 个文件：")
        for f in generated_files:
            print(f"  - {f}")
    else:
        print("\n未生成任何文件")


if __name__ == '__main__':
    main()
